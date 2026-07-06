from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def markdown_to_docx_bytes(markdown: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(58)
    section.bottom_margin = Pt(58)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
            continue
        if line.startswith("# "):
            paragraph = document.add_heading(line[2:], level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
            continue
        document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
